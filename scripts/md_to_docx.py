#!/usr/bin/env python3
"""Convert AutoResection.md to AutoResection.docx using python-docx."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


def add_code_block(doc: Document, lines: list[str]) -> None:
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        if i < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(ncols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = row[c_idx] if c_idx < len(row) else ""


def convert(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].startswith("|"):
                if re.match(r"^\|\s*[-:| ]+\|\s*$", lines[i]):
                    i += 1
                    continue
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                table_rows.append(cells)
                i += 1
            add_table(doc, table_rows)
            doc.add_paragraph()
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        # Inline formatting: strip markdown links to text (url)
        cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        cleaned = re.sub(r"\*(.+?)\*", r"\1", cleaned)
        cleaned = re.sub(r"`(.+?)`", r"\1", cleaned)
        cleaned = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1 (\2)", cleaned)
        doc.add_paragraph(cleaned)
        i += 1

    doc.save(docx_path)


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "AutoResection.md"
    out = root / "AutoResection.docx"
    if len(sys.argv) >= 2:
        md = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        out = Path(sys.argv[2])
    convert(md, out)
    print(f"Wrote {out}")
